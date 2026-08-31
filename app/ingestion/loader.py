"""文档加载：按扩展名解析文本、Markdown、PDF、Word 文件。"""
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path

SUPPORTED_EXTENSIONS = {".txt", ".md", ".markdown", ".pdf", ".docx"}


class UnsupportedFormatError(ValueError):
    """不支持的文档格式。"""


@dataclass
class Document:
    """解析后的文档：正文文本 + 元数据。"""

    text: str
    metadata: dict = field(default_factory=dict)


def load_document(filename: str, content: bytes) -> Document:
    """按文件扩展名解析二进制内容为 Document。

    Args:
        filename: 原始文件名（用于判断格式与记录来源）。
        content: 文件二进制内容。
    """
    ext = Path(filename).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFormatError(
            f"不支持的文件格式: {ext}，支持: {sorted(SUPPORTED_EXTENSIONS)}"
        )

    text = _parse(filename, ext, content)
    return Document(text=text, metadata={"source": filename})


def _parse(filename: str, ext: str, content: bytes) -> str:
    if ext in {".txt", ".md", ".markdown"}:
        return content.decode("utf-8", errors="replace")
    if ext == ".pdf":
        return _parse_pdf(content)
    if ext == ".docx":
        return _parse_docx(content)
    raise UnsupportedFormatError(f"不支持的文件格式: {ext}")


def _parse_pdf(content: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(content))
    pages = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        pages.append(page_text.strip())
    return "\n\n".join(pages)


def _parse_docx(content: bytes) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(BytesIO(content))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)
